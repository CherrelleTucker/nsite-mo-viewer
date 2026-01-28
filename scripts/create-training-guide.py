"""
Create MO-Viewer Training Guide as a Word Document
Comprehensive but readable format for different user types
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def create_training_guide():
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading('MO-Viewer User Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('NSITE Management Office Dashboard')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    # Version info
    version_para = doc.add_paragraph()
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_para.add_run('Version 1.0 | January 2026').font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

    # ===================
    # TABLE OF CONTENTS
    # ===================
    doc.add_heading('Table of Contents', level=1)

    toc_items = [
        ('1. Getting Started', 'All Users'),
        ('2. Implementation Team Guide', 'Solution Lifecycle'),
        ('3. SEP Team Guide', 'Stakeholder Engagement'),
        ('4. Comms Team Guide', 'Stories & Outreach'),
        ('5. Team Member Guide', 'Daily Operations'),
        ('6. Admin Guide', 'Data & Configuration'),
        ('7. Quick Reference', 'Tips & Shortcuts'),
    ]

    for item, desc in toc_items:
        p = doc.add_paragraph()
        p.add_run(item).bold = True
        p.add_run(f'  —  {desc}')

    doc.add_page_break()

    # ===================
    # SECTION 1: GETTING STARTED
    # ===================
    doc.add_heading('1. Getting Started', level=1)

    doc.add_heading('What is MO-Viewer?', level=2)
    doc.add_paragraph(
        'MO-Viewer is the central dashboard for the NSITE Management Office. '
        'It brings together information from across the MO into a single, searchable interface.'
    )

    doc.add_paragraph('The platform provides visibility into:')
    bullets = [
        'Solutions across their lifecycle (Preformulation → Closeout)',
        'Stakeholder engagement through the SEP pipeline',
        'Communications including stories and events',
        'Team operations including meetings and availability'
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')

    doc.add_heading('Logging In', level=2)
    steps = [
        'Navigate to the MO-Viewer URL',
        'Enter your NASA email address',
        'Enter the team passphrase',
        'You\'ll be redirected to the dashboard'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_heading('Navigation Overview', level=2)
    doc.add_paragraph('The top navigation bar provides access to all sections:')

    nav_table = doc.add_table(rows=9, cols=3)
    nav_table.style = 'Table Grid'

    nav_data = [
        ('Tab', 'Purpose', 'Primary Users'),
        ('Implementation', 'Solution lifecycle tracking', 'Implementation'),
        ('SEP', 'Stakeholder engagement pipeline', 'SEP'),
        ('Comms', 'Stories, events, messaging', 'Comms'),
        ('Contacts', 'Stakeholder directory', 'All'),
        ('Schedule', 'Milestone timeline', 'All'),
        ('Team', 'Internal operations', 'All'),
        ('Reports', 'Analytics and exports', 'All'),
        ('About', 'Platform documentation', 'All'),
    ]

    for i, row_data in enumerate(nav_data):
        row = nav_table.rows[i]
        for j, cell_text in enumerate(row_data):
            row.cells[j].text = cell_text
            if i == 0:
                row.cells[j].paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # ===================
    # SECTION 2: IMPLEMENTATION
    # ===================
    doc.add_heading('2. Implementation Team Guide', level=1)

    doc.add_paragraph(
        'The Implementation page is your home base for tracking solutions across their lifecycle.'
    )

    doc.add_heading('Solution Cards', level=2)
    doc.add_paragraph(
        'Each solution is displayed as a card showing key status information. '
        'Cards are color-coded by lifecycle phase and show the solution name, cycle, and current status.'
    )

    doc.add_paragraph('Click any card to see full details:')
    details = [
        'Official name and alternate names',
        'Current lifecycle phase and cycle',
        'Assigned DAAC',
        'Team contacts (Lead, RA Rep, EA Advocate)',
        'Decision gate dates (ATP, F2I, ORR, Closeout)',
        'Key documents with direct links',
        'Recent updates from meeting notes'
    ]
    for item in details:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Filtering Solutions', level=2)
    doc.add_paragraph('Use filters to narrow down what you see:')
    filters = [
        ('Phase Filter', 'Preformulation, Formulation, Implementation, Operations, Closeout'),
        ('Cycle Filter', 'C1 (2016), C2 (2018), C3 (2020), C4 (2022), C5 (2024), C6 (2026)'),
        ('Search', 'Type any keyword to find matching solutions'),
    ]
    for name, desc in filters:
        p = doc.add_paragraph()
        p.add_run(f'{name}: ').bold = True
        p.add_run(desc)

    doc.add_heading('Schedule View', level=2)
    doc.add_paragraph(
        'The Schedule page shows a visual timeline of solution milestones. '
        'This is useful for seeing upcoming decision gates across all solutions at once.'
    )

    doc.add_heading('Common Tasks', level=2)

    doc.add_paragraph().add_run('Finding a Solution\'s Status:').bold = True
    doc.add_paragraph('1. Go to Implementation page')
    doc.add_paragraph('2. Use search or filters to find the solution')
    doc.add_paragraph('3. Click the card to view details')
    doc.add_paragraph('4. Check "Recent Updates" for latest information')

    doc.add_paragraph()
    doc.add_paragraph().add_run('Checking Decision Gate Dates:').bold = True
    doc.add_paragraph('• Schedule page shows timeline view across all solutions')
    doc.add_paragraph('• Or: Implementation → click solution → Milestones section')

    doc.add_page_break()

    # ===================
    # SECTION 3: SEP
    # ===================
    doc.add_heading('3. SEP Team Guide', level=1)

    doc.add_paragraph(
        'The SEP page is your command center for stakeholder engagement. '
        'It shows solutions organized by their progress through the SEP pipeline.'
    )

    doc.add_heading('The SEP Pipeline', level=2)
    doc.add_paragraph('Solutions progress through working sessions (WS) and touchpoints (TP):')

    pipeline = doc.add_paragraph()
    pipeline.add_run('WS1 → TP4 → WS2 → TP5 → WS3 → TP6 → WS4 → TP7 → WS5 → TP8').bold = True
    pipeline.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    tp_table = doc.add_table(rows=6, cols=2)
    tp_table.style = 'Table Grid'
    tp_data = [
        ('Milestone', 'Description'),
        ('WS1-WS5', 'Working Sessions: collaborative design workshops'),
        ('TP4', 'Invitation to CoDesign'),
        ('TP5-TP6', 'Design Review and Beta Testing'),
        ('TP7', 'Soft Launch & Training'),
        ('TP8', 'Closeout Stories'),
    ]
    for i, row_data in enumerate(tp_data):
        row = tp_table.rows[i]
        for j, cell_text in enumerate(row_data):
            row.cells[j].text = cell_text
            if i == 0:
                row.cells[j].paragraphs[0].runs[0].bold = True

    doc.add_heading('Solution Details', level=2)
    doc.add_paragraph('Click any solution in the pipeline to see:')
    details = [
        'Current SEP milestone status',
        'Full stakeholder list with contact info',
        'Engagement history',
        'Quick actions: Email stakeholders, Log engagement'
    ]
    for item in details:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Working with Stakeholders', level=2)

    doc.add_paragraph().add_run('Emailing Stakeholders:').bold = True
    doc.add_paragraph('1. Click a solution → Stakeholders section')
    doc.add_paragraph('2. Click "Email All" for group email, or click individual contact')
    doc.add_paragraph('3. Use Template dropdown for pre-written touchpoint/session emails')
    doc.add_paragraph('4. Click "Log as Engagement" after sending to record the interaction')

    doc.add_paragraph()
    doc.add_paragraph().add_run('Logging Engagements:').bold = True
    doc.add_paragraph('After any stakeholder interaction:')
    doc.add_paragraph('1. Click "Log as Engagement" in the email composer, OR')
    doc.add_paragraph('2. Go to Contacts → find stakeholder → Add Engagement')
    doc.add_paragraph('3. Select activity type (Email, Meeting, Call)')
    doc.add_paragraph('4. Add notes about the interaction')

    doc.add_heading('Contacts Directory', level=2)
    doc.add_paragraph(
        'The Contacts page is your stakeholder directory. Search by name, email, agency, '
        'or department. Filter by solution or survey year to find specific groups.'
    )

    doc.add_page_break()

    # ===================
    # SECTION 4: COMMS
    # ===================
    doc.add_heading('4. Comms Team Guide', level=1)

    doc.add_paragraph(
        'The Comms page is your hub for tracking stories, events, and outreach activities.'
    )

    doc.add_heading('Stories Pipeline', level=2)
    doc.add_paragraph('Track stories from concept to publication:')

    stages = doc.add_paragraph()
    stages.add_run('Idea → Draft → Reviewed → Submitted → Published').bold = True
    stages.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph('Filter stories by solution, channel, or status to manage your pipeline.')

    doc.add_heading('Events', level=2)
    doc.add_paragraph('Track conferences and outreach events:')
    items = [
        'View upcoming events',
        'Manage guest lists (linked to Contacts)',
        'Add prep notes and materials',
        'Record post-event outcomes'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Finding Content for Solutions', level=2)
    doc.add_paragraph('Need messaging for a solution? Go to Implementation → click solution → Communications section:')
    items = [
        'Key messages (approved talking points)',
        'Thematic areas',
        'Focus type (Climate, Disasters, etc.)',
        'Links to existing public content'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ===================
    # SECTION 5: TEAM MEMBER
    # ===================
    doc.add_heading('5. Team Member Guide', level=1)

    doc.add_paragraph(
        'The Team page is your home for internal operations. Use the toggle buttons to switch between views.'
    )

    doc.add_heading('Overview', level=2)
    doc.add_paragraph('Two sections:')

    doc.add_paragraph().add_run('Team Profiles: ').bold = True
    doc.add_paragraph('See all team members. Click a profile for details. Filter by team.')

    doc.add_paragraph().add_run('Office Availability: ').bold = True
    doc.add_paragraph('See who\'s out—vacation, travel, 9/80 days, holidays. Past dates are automatically hidden.')

    doc.add_heading('Adding Your Availability', level=2)
    doc.add_paragraph('1. Team → Overview')
    doc.add_paragraph('2. Click "Add" next to Office Availability')
    doc.add_paragraph('3. Select your name')
    doc.add_paragraph('4. Choose type: Vacation, Work Travel, 9/80, Out of Office')
    doc.add_paragraph('5. Enter dates and optional notes')
    doc.add_paragraph('6. Save')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('To edit: ').italic = True
    p.add_run('Click your existing entry → update → Save or Delete')

    doc.add_heading('Actions', level=2)
    doc.add_paragraph(
        'Track action items from meetings. Toggle between "By Person" (see your tasks) '
        'and "By Status" (Kanban board view).'
    )
    doc.add_paragraph('Filter by Open, All, or Done. Click any action to update its status or reassign.')

    doc.add_heading('Kudos', level=2)
    doc.add_paragraph('Recognize teammates for great work:')
    doc.add_paragraph('1. Select recipient')
    doc.add_paragraph('2. Choose category (Teamwork, Innovation, Above & Beyond, Mentorship, Delivery)')
    doc.add_paragraph('3. Write a message (280 characters max)')
    doc.add_paragraph('4. Optionally post to Slack')
    doc.add_paragraph('5. Submit')

    doc.add_heading('Meetings', level=2)
    doc.add_paragraph(
        'Find recurring meeting schedules with links to video calls, agendas, and notes. '
        'Organized by day of the week.'
    )

    doc.add_heading('Documents', level=2)
    doc.add_paragraph('Quick access to directing documents: Project Plans, Style Guide, Process Docs, etc.')

    doc.add_heading('Parking Lot', level=2)
    doc.add_paragraph('Capture ideas and topics for future discussion. Add new items and track status.')

    doc.add_page_break()

    # ===================
    # SECTION 6: ADMIN
    # ===================
    doc.add_heading('6. Admin Guide', level=1)

    doc.add_paragraph(
        'This section is for dashboard administrators who manage data and configuration.'
    )

    doc.add_heading('Architecture Overview', level=2)

    arch_para = doc.add_paragraph()
    arch_para.add_run(
        'Source Documents → Sync Scripts → Databases → API Library → Web App'
    ).bold = True
    arch_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph(
        'Data flows from Google Docs (meeting notes, agendas) through sync scripts '
        'into Google Sheets databases, which are queried by the web app.'
    )

    doc.add_heading('Key Databases', level=2)

    db_table = doc.add_table(rows=7, cols=2)
    db_table.style = 'Table Grid'
    db_data = [
        ('Database', 'Purpose'),
        ('MO-DB_Solutions', 'Solution master data (48 solutions)'),
        ('MO-DB_Contacts', 'Stakeholder contacts (4,200+ records)'),
        ('MO-DB_Updates', 'Meeting note updates by year'),
        ('MO-DB_Engagements', 'Stakeholder interaction logs'),
        ('MO-DB_Availability', 'Team availability calendar'),
        ('MO-DB_Config', 'Configuration keys and settings'),
    ]
    for i, row_data in enumerate(db_data):
        row = db_table.rows[i]
        for j, cell_text in enumerate(row_data):
            row.cells[j].text = cell_text
            if i == 0:
                row.cells[j].paragraphs[0].runs[0].bold = True

    doc.add_heading('Common Admin Tasks', level=2)

    doc.add_paragraph().add_run('Adding a New User:').bold = True
    doc.add_paragraph('1. Open MO-DB_Access spreadsheet')
    doc.add_paragraph('2. Add row with user\'s email')
    doc.add_paragraph('3. User can now log in with team passphrase')

    doc.add_paragraph()
    doc.add_paragraph().add_run('Adding a New Solution:').bold = True
    doc.add_paragraph('1. Open MO-DB_Solutions')
    doc.add_paragraph('2. Add new row with: core_id, core_official_name, core_cycle, admin_lifecycle_phase')
    doc.add_paragraph('3. Fill in other fields as available')
    doc.add_paragraph('4. Solution appears immediately on next page load')

    doc.add_paragraph()
    doc.add_paragraph().add_run('Running Sync Scripts:').bold = True
    doc.add_paragraph('1. Open the relevant database spreadsheet')
    doc.add_paragraph('2. Go to Extensions → Apps Script')
    doc.add_paragraph('3. Run the appropriate sync function')
    doc.add_paragraph('4. Check execution logs for errors')

    doc.add_heading('Troubleshooting', level=2)

    issues = [
        ('Data not appearing', 'Check source sheet, verify column headers match exactly, clear browser cache'),
        ('User can\'t log in', 'Verify email in MO-DB_Access, confirm passphrase'),
        ('Slow performance', 'Check for large data responses, review sync script efficiency'),
    ]
    for issue, solution in issues:
        p = doc.add_paragraph()
        p.add_run(f'{issue}: ').bold = True
        p.add_run(solution)

    doc.add_page_break()

    # ===================
    # SECTION 7: QUICK REFERENCE
    # ===================
    doc.add_heading('7. Quick Reference', level=1)

    doc.add_heading('Tips for Everyone', level=2)
    tips = [
        'Use the Refresh button (top right) to reload data',
        'Search works on most pages—use it liberally',
        'Reset filters if something seems missing',
        'Click the About tab for data source documentation',
        'Press Escape to close modals and drawers'
    ]
    for tip in tips:
        doc.add_paragraph(tip, style='List Bullet')

    doc.add_heading('Common Filters', level=2)

    doc.add_paragraph().add_run('Solutions:').bold = True
    doc.add_paragraph('• By Cycle: C1, C2, C3, C4, C5, C6')
    doc.add_paragraph('• By Phase: Preformulation, Formulation, Implementation, Operations, Closeout')

    doc.add_paragraph().add_run('Contacts:').bold = True
    doc.add_paragraph('• By Role: Primary SME, Secondary SME, Survey Submitter')
    doc.add_paragraph('• By Survey Year: 2016, 2018, 2020, 2022, 2024')

    doc.add_paragraph().add_run('Actions:').bold = True
    doc.add_paragraph('• By Status: Open, All, Done')
    doc.add_paragraph('• By Source: MO (Internal), SEP')

    doc.add_heading('Getting Help', level=2)
    doc.add_paragraph('• Check the About page for platform documentation')
    doc.add_paragraph('• Contact your team lead for access issues')
    doc.add_paragraph('• Contact the dashboard admin for data issues')

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('— End of Guide —').italic = True

    # Save
    output_path = os.path.join(os.path.dirname(__file__), '..', 'MO-Viewer-Training-Guide.docx')
    doc.save(output_path)
    print(f'Training guide saved to: {output_path}')
    return output_path

if __name__ == '__main__':
    create_training_guide()
