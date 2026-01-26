# -*- coding: utf-8 -*-
"""Explore structure of older Monthly meeting Word documents"""

import docx
from docx.oxml.ns import qn
from pathlib import Path

# Sample files from different fiscal years
sample_files = [
    Path(r"C:\Users\cjtucke3\Documents\Personal\MO-development\source-archives\Monthly Project Status Updates\FY21\2021-05-10 SNWG Monthly Meeting.docx"),
    Path(r"C:\Users\cjtucke3\Documents\Personal\MO-development\source-archives\Monthly Project Status Updates\FY22\2022-03-28 SNWG Monthly Meeting.docx"),
    Path(r"C:\Users\cjtucke3\Documents\Personal\MO-development\source-archives\Monthly Project Status Updates\FY23\2023-01-23 SNWG Monthly Meeting.docx"),
]

for doc_path in sample_files:
    if not doc_path.exists():
        print(f"File not found: {doc_path}")
        continue

    print(f"\n{'='*60}")
    print(f"FILE: {doc_path.name}")
    print(f"{'='*60}")

    try:
        doc = docx.Document(doc_path)
    except Exception as e:
        print(f"Error opening: {e}")
        continue

    # Show paragraphs
    print(f"\nParagraphs ({len(doc.paragraphs)} total):")
    for i, para in enumerate(doc.paragraphs[:30]):
        text = para.text.strip()
        if text:
            style = para.style.name if para.style else "None"
            print(f"  [{i}] ({style}) {text[:80]}...")

    # Show tables
    print(f"\nTables ({len(doc.tables)} total):")
    for t_idx, table in enumerate(doc.tables[:3]):
        print(f"\n  Table {t_idx}: {len(table.rows)} rows x {len(table.columns)} cols")
        for r_idx, row in enumerate(table.rows[:5]):
            cells_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()[:40]
                if cell_text:
                    cells_text.append(cell_text)
            if cells_text:
                print(f"    Row {r_idx}: {' | '.join(cells_text)}")
