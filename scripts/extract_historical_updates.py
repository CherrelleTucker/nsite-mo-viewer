# -*- coding: utf-8 -*-
"""
Extract Historical Updates from Weekly Internal Planning Meeting Documents
===========================================================================
Parses .docx files across multiple fiscal years with varying formats.
Maps solution names to core_ids using MO-DB_Solutions database.
Outputs a CSV ready for import into MO-DB_Updates.

Usage: python extract_historical_updates.py
"""

import docx
from docx.oxml.ns import qn
from pathlib import Path
import pandas as pd
import csv
import re
from datetime import datetime
import uuid

# Configuration
NEW_MARKER = '🆕'
BASE_PATH = Path(r"C:\Users\cjtucke3\Documents\Personal\MO-development\source-archives\Weekly Internal Planning")
SOLUTIONS_PATH = Path(r"C:\Users\cjtucke3\Documents\Personal\MO-development\nsite-mo-viewer\database-files\MO-Viewer Databases\MO-DB_Solutions.xlsx")
OUTPUT_PATH = Path(r"C:\Users\cjtucke3\Documents\Personal\MO-development\nsite-mo-viewer\database-files\historical_updates_import.csv")

# CSV columns matching MO-DB_Updates
CSV_HEADERS = [
    'update_id', 'solution_id', 'update_text', 'source_document',
    'source_category', 'source_url', 'source_tab', 'meeting_date',
    'created_at', 'created_by'
]


def generate_update_id():
    """Generate a unique update ID"""
    date_str = datetime.now().strftime('%Y%m%d')
    random_part = str(uuid.uuid4())[:4].upper()
    return f"UPD_{date_str}_{random_part}"


def build_solution_mapping():
    """Build mapping from solution names/aliases to core_ids"""
    df = pd.read_excel(SOLUTIONS_PATH)
    mapping = {}

    for _, row in df.iterrows():
        core_id = str(row.get('core_id', '')).strip()
        if not core_id or core_id == 'nan':
            continue

        official_name = str(row.get('core_official_name', '')).strip()
        alternates = str(row.get('core_alternate_names', '')).strip()

        # Add official name
        if official_name and official_name != 'nan':
            mapping[official_name.lower()] = core_id
            # Also add without parenthetical
            clean = re.sub(r'\s*\([^)]+\)\s*', '', official_name).strip()
            if clean:
                mapping[clean.lower()] = core_id

        # Add core_id itself as a match
        mapping[core_id.lower()] = core_id

        # Add alternates (pipe or comma separated)
        if alternates and alternates != 'nan':
            for alt in re.split(r'[|,]', alternates):
                alt = alt.strip()
                if alt:
                    mapping[alt.lower()] = core_id

    return mapping


def find_core_id(text, solution_mapping):
    """Find core_id for a solution name using the mapping"""
    if not text:
        return None

    text_clean = text.strip().lower()

    # Direct match
    if text_clean in solution_mapping:
        return solution_mapping[text_clean]

    # Try without parenthetical suffix
    no_paren = re.sub(r'\s*\([^)]+\)\s*$', '', text_clean).strip()
    if no_paren in solution_mapping:
        return solution_mapping[no_paren]

    # Try partial matches for known patterns
    for name, core_id in solution_mapping.items():
        if name in text_clean or text_clean in name:
            return core_id

    return None


def extract_hyperlinks_from_paragraph(paragraph):
    """Extract hyperlinks from a paragraph"""
    hyperlinks = []

    for element in paragraph._element.iter():
        if element.tag == qn('w:hyperlink'):
            r_id = element.get(qn('r:id'))
            if r_id:
                try:
                    rel = paragraph.part.rels[r_id]
                    url = rel.target_ref
                    text_parts = []
                    for t_elem in element.iter(qn('w:t')):
                        if t_elem.text:
                            text_parts.append(t_elem.text)
                    text = ''.join(text_parts)
                    if text and url:
                        hyperlinks.append((text, url))
                except:
                    pass
    return hyperlinks


def get_list_items_from_cell(cell):
    """Extract list items with their nesting levels from a cell"""
    items = []

    for paragraph in cell.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        # Determine nesting level
        nesting = 0
        pPr = paragraph._element.find(qn('w:pPr'))
        if pPr is not None:
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                ilvl = numPr.find(qn('w:ilvl'))
                if ilvl is not None:
                    nesting = int(ilvl.get(qn('w:val'), 0))

            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                left = ind.get(qn('w:left'))
                if left:
                    nesting = max(nesting, int(int(left) / 720))

        # Extract hyperlinks
        hyperlinks = extract_hyperlinks_from_paragraph(paragraph)
        para_text = text
        for link_text, url in hyperlinks:
            if link_text in para_text:
                para_text = para_text.replace(link_text, f"[{link_text}]({url})")

        items.append({
            'text': para_text,
            'nesting': nesting
        })

    return items


def extract_solution_name(text):
    """Extract solution name from text, handling various formats"""
    if not text:
        return ''

    text = text.strip()

    # Pattern: "Name [core_id]" - extract core_id
    bracket_match = re.match(r'^(.+?)\s*\[([^\]]+)\]$', text)
    if bracket_match:
        return bracket_match.group(2).strip()

    # Pattern: "[core_id]" only
    only_bracket = re.match(r'^\[([^\]]+)\]$', text)
    if only_bracket:
        return only_bracket.group(1).strip()

    # Pattern: "Name (Provider)" - return full name for mapping
    return text


def is_solution_header(text, solution_mapping):
    """Check if text looks like a solution header - conservative matching"""
    if not text:
        return False

    text_clean = text.strip()

    # Has [core_id] suffix - definitely a solution
    if re.search(r'\[[^\]]+\]$', text_clean):
        return True

    # Known solution name from database - this is the primary way to identify
    if find_core_id(text_clean, solution_mapping):
        return True

    # Don't match generic administrative headers
    # Only strict database matches count
    return False


def parse_items_notes_format(doc, doc_path, meeting_date, solution_mapping):
    """Parse FY24+ style documents with Items/Notes table format"""
    updates = []

    for table in doc.tables:
        # Check if this is the main content table
        if len(table.rows) < 2:
            continue

        header_text = table.rows[0].cells[0].text.lower() if table.rows[0].cells else ''
        if 'items' not in header_text and 'notes' not in header_text:
            # Try second row for content tables that don't have header
            pass

        for row in table.rows:
            if len(row.cells) < 2:
                continue

            cell = row.cells[1]  # Notes column
            items = get_list_items_from_cell(cell)

            current_solution = None
            current_solution_name = None
            current_solution_nesting = -1

            for i, item in enumerate(items):
                text = item['text']
                nesting = item['nesting']

                if not text:
                    continue

                # Check if this is a solution header
                if is_solution_header(text, solution_mapping):
                    solution_name = extract_solution_name(text)
                    core_id = find_core_id(solution_name, solution_mapping)

                    if nesting <= current_solution_nesting or current_solution is None:
                        current_solution = core_id if core_id else solution_name
                        current_solution_name = solution_name
                        current_solution_nesting = nesting
                    continue

                # This is content under a solution
                if current_solution:
                    # Check for 🆕 marker (explicit new update)
                    has_new_marker = NEW_MARKER in text

                    # Clean up text
                    update_text = text.replace(NEW_MARKER, '').strip()

                    # Skip empty or very short updates
                    if len(update_text) < 10:
                        continue

                    # Skip items that are just dates or links
                    if re.match(r'^[\d\-/]+$', update_text):
                        continue

                    # Skip action items (captured in separate database)
                    # Filter any update containing action: or action item
                    if re.search(r'action[:\s]', update_text, re.IGNORECASE):
                        continue
                    if 'action item' in update_text.lower():
                        continue

                    # Collect child items (skip action items)
                    child_parts = [update_text]
                    for j in range(i + 1, len(items)):
                        child = items[j]
                        if child['nesting'] <= nesting:
                            break
                        child_text = child['text']
                        if child_text and NEW_MARKER not in child_text:
                            # Skip action items in children
                            if re.match(r'^action[:\s]', child_text, re.IGNORECASE):
                                continue
                            indent = '  ' * (child['nesting'] - nesting)
                            child_parts.append(f"{indent}• {child_text}")

                    # Filter out action lines from final text
                    filtered_parts = []
                    for part in child_parts:
                        # Skip lines containing action items
                        if re.search(r'action[:\s]', part, re.IGNORECASE):
                            continue
                        filtered_parts.append(part)

                    if not filtered_parts:
                        continue

                    full_text = '\n'.join(filtered_parts)

                    # Skip if nothing left after filtering
                    if len(full_text.strip()) < 10:
                        continue

                    updates.append({
                        'update_id': generate_update_id(),
                        'solution_id': current_solution,
                        'update_text': full_text,
                        'source_document': 'Internal Planning',
                        'source_category': 'MO',
                        'source_url': '',
                        'source_tab': doc_path.name,
                        'meeting_date': meeting_date,
                        'created_at': datetime.now().isoformat(),
                        'created_by': 'historical_import',
                        'has_new_marker': has_new_marker
                    })

    return updates


def parse_agenda_format(doc, doc_path, meeting_date, solution_mapping):
    """Parse FY22-FY23 style documents with agenda format"""
    updates = []

    # These documents often have content after the header table
    # Look for bullet lists in paragraphs or later tables

    current_solution = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Check if this is a solution header
        if is_solution_header(text, solution_mapping):
            solution_name = extract_solution_name(text)
            core_id = find_core_id(solution_name, solution_mapping)
            current_solution = core_id if core_id else solution_name
            continue

        # Content under solution
        if current_solution and len(text) > 15:
            # Skip action items (captured in separate database)
            if re.search(r'action[:\s]', text, re.IGNORECASE):
                continue
            if 'action item' in text.lower():
                continue
            updates.append({
                'update_id': generate_update_id(),
                'solution_id': current_solution,
                'update_text': text,
                'source_document': 'Internal Planning',
                'source_category': 'MO',
                'source_url': '',
                'source_tab': doc_path.name,
                'meeting_date': meeting_date,
                'created_at': datetime.now().isoformat(),
                'created_by': 'historical_import',
                'has_new_marker': False
            })

    # Also check tables beyond the header
    for table in doc.tables[1:] if len(doc.tables) > 1 else []:
        for row in table.rows:
            for cell in row.cells:
                items = get_list_items_from_cell(cell)
                for item in items:
                    text = item['text']
                    if not text or len(text) < 15:
                        continue

                    if is_solution_header(text, solution_mapping):
                        solution_name = extract_solution_name(text)
                        core_id = find_core_id(solution_name, solution_mapping)
                        current_solution = core_id if core_id else solution_name
                    elif current_solution:
                        # Skip action items (captured in separate database)
                        if re.search(r'action[:\s]', text, re.IGNORECASE):
                            continue
                        if 'action item' in text.lower():
                            continue
                        updates.append({
                            'update_id': generate_update_id(),
                            'solution_id': current_solution,
                            'update_text': text,
                            'source_document': 'Internal Planning',
                            'source_category': 'MO',
                            'source_url': '',
                            'source_tab': doc_path.name,
                            'meeting_date': meeting_date,
                            'created_at': datetime.now().isoformat(),
                            'created_by': 'historical_import',
                            'has_new_marker': False
                        })

    return updates


def detect_document_format(doc):
    """Detect which format a document uses"""
    if not doc.tables:
        return 'no_tables'

    first_table = doc.tables[0]
    if not first_table.rows:
        return 'unknown'

    first_cell = first_table.rows[0].cells[0].text.lower() if first_table.rows[0].cells else ''

    if 'items' in first_cell or 'notes' in first_cell:
        return 'items_notes'
    elif 'location' in first_cell or 'meeting' in first_cell or 'date' in first_cell:
        return 'agenda'
    else:
        return 'items_notes'  # Default to items/notes parsing


def parse_document(doc_path, meeting_date, solution_mapping):
    """Parse a document and extract updates"""
    try:
        doc = docx.Document(doc_path)
    except Exception as e:
        print(f"  Error opening {doc_path.name}: {e}")
        return []

    doc_format = detect_document_format(doc)

    if doc_format == 'items_notes':
        return parse_items_notes_format(doc, doc_path, meeting_date, solution_mapping)
    elif doc_format == 'agenda':
        return parse_agenda_format(doc, doc_path, meeting_date, solution_mapping)
    else:
        return []


def extract_date_from_filename(filename):
    """Extract meeting date from filename"""
    match = re.match(r'(\d{4}-\d{2}-\d{2})', filename)
    if match:
        return match.group(1)
    return None


def parse_consolidated_document(doc_path, solution_mapping):
    """
    Parse consolidated document (e.g., Weekly Internal Planning Meeting_NSITE MO_C0_2026.docx)
    These files contain multiple meetings with date markers in MM_DD format.
    Iterates through document body elements in order to correctly associate tables with dates.
    """
    try:
        doc = docx.Document(doc_path)
    except Exception as e:
        print(f"  Error opening {doc_path.name}: {e}")
        return []

    # Extract year from filename (e.g., _C0_2026.docx -> 2026)
    year_match = re.search(r'_C0_(\d{4})\.docx', doc_path.name)
    current_year = int(year_match.group(1)) if year_match else datetime.now().year

    updates = []
    current_meeting_date = None
    date_pattern = re.compile(r'^(\d{2})_(\d{2})$')  # MM_DD pattern
    year_pattern_re = re.compile(r'^(20\d{2})$')  # Year pattern
    meeting_count = 0

    # Build a list of (table_element, meeting_date) by iterating body in order
    table_dates = []

    for element in doc.element.body:
        if element.tag == qn('w:p'):  # Paragraph
            text = ''.join(t.text or '' for t in element.iter(qn('w:t'))).strip()
            if not text:
                continue

            # Check for year marker
            year_match = year_pattern_re.match(text)
            if year_match:
                current_year = int(year_match.group(1))
                continue

            # Check for date marker MM_DD
            date_match = date_pattern.match(text)
            if date_match:
                month = date_match.group(1)
                day = date_match.group(2)
                current_meeting_date = f"{current_year}-{month}-{day}"
                meeting_count += 1

        elif element.tag == qn('w:tbl'):  # Table
            table_dates.append((element, current_meeting_date))

    print(f"    Found {meeting_count} meeting dates, {len(table_dates)} tables")

    # Now process tables with their assigned dates
    for table_element, meeting_date in table_dates:
        # Convert element back to Table object
        table = docx.table.Table(table_element, doc)

        if len(table.rows) < 2:
            continue

        # Check if this is a content table
        header_text = table.rows[0].cells[0].text.lower() if table.rows[0].cells else ''

        for row in table.rows:
            if len(row.cells) < 2:
                continue

            cell = row.cells[1]  # Notes column (second column)
            items = get_list_items_from_cell(cell)

            current_solution = None
            current_solution_nesting = -1

            for item_idx, item in enumerate(items):
                text = item['text']
                nesting = item['nesting']

                if not text:
                    continue

                # Check if this is a solution header
                if is_solution_header(text, solution_mapping):
                    solution_name = extract_solution_name(text)
                    core_id = find_core_id(solution_name, solution_mapping)
                    if nesting <= current_solution_nesting or current_solution is None:
                        current_solution = core_id if core_id else solution_name
                        current_solution_nesting = nesting
                    continue

                # Content under solution
                if current_solution:
                    has_new_marker = NEW_MARKER in text
                    update_text = text.replace(NEW_MARKER, '').strip()

                    if len(update_text) < 10:
                        continue
                    if re.match(r'^[\d\-/]+$', update_text):
                        continue
                    if re.search(r'action[:\s]', update_text, re.IGNORECASE):
                        continue
                    if 'action item' in update_text.lower():
                        continue

                    # Get child items
                    child_parts = [update_text]
                    for j in range(item_idx + 1, len(items)):
                        child = items[j]
                        if child['nesting'] <= nesting:
                            break
                        child_text = child['text']
                        if child_text and NEW_MARKER not in child_text:
                            if re.search(r'action[:\s]', child_text, re.IGNORECASE):
                                continue
                            indent = '  ' * (child['nesting'] - nesting)
                            child_parts.append(f"{indent}• {child_text}")

                    full_text = '\n'.join(child_parts)
                    if len(full_text.strip()) < 10:
                        continue

                    updates.append({
                        'update_id': generate_update_id(),
                        'solution_id': current_solution,
                        'update_text': full_text,
                        'source_document': 'Internal Planning',
                        'source_category': 'MO',
                        'source_url': '',
                        'source_tab': meeting_date or doc_path.name,
                        'meeting_date': meeting_date or '',
                        'created_at': datetime.now().isoformat(),
                        'created_by': 'historical_import',
                        'has_new_marker': has_new_marker
                    })

    return updates


def main():
    all_updates = []
    files_processed = 0

    print("Building solution name to core_id mapping...")
    solution_mapping = build_solution_mapping()
    print(f"  Loaded {len(solution_mapping)} name/alias mappings")
    print()

    print("Extracting historical updates from Weekly Internal Planning documents...")
    print(f"Base path: {BASE_PATH}")
    print()

    # Process FY folders
    for fy_folder in sorted(BASE_PATH.glob("FY*")):
        if fy_folder.is_dir():
            print(f"Processing {fy_folder.name}...")
            fy_updates = 0
            for doc_file in sorted(fy_folder.glob("*.docx")):
                meeting_date = extract_date_from_filename(doc_file.name)
                if meeting_date:
                    updates = parse_document(doc_file, meeting_date, solution_mapping)
                    if updates:
                        fy_updates += len(updates)
                        all_updates.extend(updates)
                    files_processed += 1
            print(f"  {fy_updates} updates from {fy_folder.name}")

    # Process consolidated docs (_C0_ files for 2025, 2026)
    for doc_file in sorted(BASE_PATH.glob("*_C0_*.docx")):
        print(f"Processing consolidated file: {doc_file.name}...")
        updates = parse_consolidated_document(doc_file, solution_mapping)
        if updates:
            print(f"  Found {len(updates)} updates")
            all_updates.extend(updates)
        files_processed += 1

    print()
    print(f"Total files processed: {files_processed}")
    print(f"Total updates found: {len(all_updates)}")

    # Count updates with explicit NEW markers
    new_marked = sum(1 for u in all_updates if u.get('has_new_marker'))
    print(f"Updates with NEW marker: {new_marked}")

    # Count by solution
    by_solution = {}
    for u in all_updates:
        sol = u['solution_id']
        by_solution[sol] = by_solution.get(sol, 0) + 1

    print(f"Unique solutions: {len(by_solution)}")
    print("\nTop 10 solutions by update count:")
    for sol, count in sorted(by_solution.items(), key=lambda x: -x[1])[:10]:
        print(f"  {sol}: {count}")

    # Remove internal tracking field before export
    for u in all_updates:
        u.pop('has_new_marker', None)

    # Write CSV
    if all_updates:
        OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
        with open(OUTPUT_PATH, 'w', newline='', encoding='utf-8') as f:
            writer = csv.DictWriter(f, fieldnames=CSV_HEADERS)
            writer.writeheader()
            writer.writerows(all_updates)
        print(f"\nCSV written to: {OUTPUT_PATH}")
    else:
        print("\nNo updates found to export.")


if __name__ == '__main__':
    main()
