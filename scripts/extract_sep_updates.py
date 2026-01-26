# -*- coding: utf-8 -*-
"""
Extract SEP (Stakeholder Engagement Program) Updates from Word Documents
=========================================================================
Parses .docx files from SEP folder structure.
Maps filenames to Google Drive URLs using file log.
Outputs consolidated Excel file for import to MO-DB_Updates.

Usage: python extract_sep_updates.py
"""

import docx
from docx.oxml.ns import qn
from pathlib import Path
import pandas as pd
import csv
import re
from datetime import datetime
import uuid

# Configuration
BASE_PATH = Path(r"C:\Users\cjtucke3\Documents\Personal\MO-development\source-archives\SEP")
FILE_LOG_PATH = Path(r"C:\Users\cjtucke3\Documents\Personal\MO-development\source-archives\file log - Sheet1.csv")
SOLUTIONS_PATH = Path(r"C:\Users\cjtucke3\Documents\Personal\MO-development\nsite-mo-viewer\database-files\MO-Viewer Databases\MO-DB_Solutions.xlsx")
OUTPUT_PATH = Path(r"C:\Users\cjtucke3\Documents\Personal\MO-development\nsite-mo-viewer\database-files\sep_updates_combined.xlsx")

# Output columns
OUTPUT_COLUMNS = [
    'update_id', 'solution_id', 'update_text', 'source_document',
    'source_category', 'source_url', 'source_tab', 'meeting_date',
    'created_at', 'created_by'
]

# Solution ID normalization
SOLUTION_ID_NORMALIZATION = {
    'hls': 'HLS',
    'opera': 'OPERA',
    'dswx': 'DSWx',
    'dist': 'DIST',
    'disp': 'DISP',
    'icesat-2': 'ICESat-2',
    'icesat2': 'ICESat-2',
    'csda': 'CSDA',
    'gaban': 'GABAN',
    'vlm': 'VLM',
    'pbl': 'PBL',
    'tempo': 'TEMPO',
    'tempo_nrt': 'TEMPO-NRT',
    'ioa': 'IoA',
    'admg': 'ADMG',
    'mwow': 'MWoW',
    'sep': 'SEP',
    'snwg': 'SNWG-MO',
    'mo': 'SNWG-MO',
    'sport': 'SPoRT',
    'servir': 'SERVIR',
    'arset': 'ARSET',
    'usgs': 'USGS',
    'usaid': 'USAID',
    'air quality': 'Air-Quality',
}

# Maximum text length
MAX_UPDATE_LENGTH = 3000


def generate_update_id():
    """Generate a unique update ID"""
    date_str = datetime.now().strftime('%Y%m%d')
    random_part = str(uuid.uuid4())[:4].upper()
    return f"UPD_{date_str}_{random_part}"


def build_url_mapping():
    """Build mapping from filename to Google Drive URL"""
    url_map = {}

    if not FILE_LOG_PATH.exists():
        print(f"Warning: File log not found at {FILE_LOG_PATH}")
        return url_map

    try:
        with open(FILE_LOG_PATH, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            next(reader)  # Skip header

            for row in reader:
                if len(row) >= 6:
                    file_title = row[0].strip()
                    doc_id = row[4].strip() if len(row) > 4 else ''
                    mime_type = row[5].strip() if len(row) > 5 else ''

                    if doc_id and file_title:
                        # Construct URL based on MIME type
                        if 'document' in mime_type:
                            url = f"https://docs.google.com/document/d/{doc_id}"
                        elif 'spreadsheet' in mime_type:
                            url = f"https://docs.google.com/spreadsheets/d/{doc_id}"
                        elif 'presentation' in mime_type:
                            url = f"https://docs.google.com/presentation/d/{doc_id}"
                        else:
                            url = f"https://drive.google.com/file/d/{doc_id}"

                        # Store by multiple key variations
                        url_map[file_title] = url
                        url_map[file_title.lower()] = url

                        # Also try without extension
                        base_name = file_title.replace('.docx', '').replace('.xlsx', '')
                        url_map[base_name] = url
                        url_map[base_name.lower()] = url

    except Exception as e:
        print(f"Error reading file log: {e}")

    return url_map


def find_url_for_file(filename, url_map):
    """Find URL for a filename using various matching strategies"""
    # Try exact match
    if filename in url_map:
        return url_map[filename]

    # Try without extension
    base = filename.replace('.docx', '').replace('.xlsx', '')
    if base in url_map:
        return url_map[base]

    # Try lowercase
    if filename.lower() in url_map:
        return url_map[filename.lower()]

    # Try partial match on date pattern
    date_match = re.match(r'(\d{4}-\d{2}-\d{2})', filename)
    if date_match:
        date_str = date_match.group(1)
        for key, url in url_map.items():
            if date_str in key:
                return url

    return ''


def normalize_solution_id(solution_id):
    """Normalize solution ID to canonical form"""
    if not solution_id:
        return solution_id
    lower = str(solution_id).lower().strip()
    return SOLUTION_ID_NORMALIZATION.get(lower, solution_id)


def extract_date_from_filename(filename):
    """Extract date from filename"""
    match = re.match(r'(\d{4}-\d{2}-\d{2})', filename)
    if match:
        return match.group(1)
    return None


def clean_update_text(text):
    """Clean boilerplate and standardize formatting"""
    if not text:
        return ''

    text = str(text)

    # Remove common boilerplate
    removals = [
        r'^\s*•\s*',
        r'^\s*○\s*',
        r'Action:\s*.*$',
        r'ACTION:\s*.*$',
    ]

    for pattern in removals:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE | re.MULTILINE)

    # Clean whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = text.strip()

    return text


def is_meaningful_update(text):
    """Check if update has real content"""
    if not text or len(text) < 30:
        return False

    text_lower = text.lower().strip()

    skip_patterns = [
        r'^n/a\s*$',
        r'^none\s*$',
        r'^no\s+update\s*$',
        r'^tbd\s*$',
    ]

    for pattern in skip_patterns:
        if re.match(pattern, text_lower):
            return False

    return True


def extract_solution_from_filename(filename):
    """Try to identify solution from filename"""
    filename_lower = filename.lower()

    # Check for known solution names in filename
    solution_patterns = [
        (r'opera', 'OPERA'),
        (r'hls', 'HLS'),
        (r'icesat', 'ICESat-2'),
        (r'csda', 'CSDA'),
        (r'sport', 'SPoRT'),
        (r'servir', 'SERVIR'),
        (r'arset', 'ARSET'),
        (r'usgs', 'USGS'),
        (r'usaid', 'USAID'),
        (r'air\s*quality', 'Air-Quality'),
        (r'co-?design', 'SEP'),
        (r'sep\s+weekly', 'SEP'),
    ]

    for pattern, solution in solution_patterns:
        if re.search(pattern, filename_lower):
            return solution

    return 'SEP'  # Default to SEP for general SEP documents


def get_paragraphs_text(doc):
    """Extract all meaningful text from document"""
    texts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text and len(text) > 10:
            texts.append(text)

    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text and len(cell_text) > 10:
                    texts.append(cell_text)

    return '\n'.join(texts)


def parse_document(doc_path, url_map):
    """Parse a Word document and extract updates"""
    try:
        doc = docx.Document(doc_path)
    except Exception as e:
        print(f"  Error opening {doc_path.name}: {e}")
        return []

    filename = doc_path.name
    meeting_date = extract_date_from_filename(filename)
    source_url = find_url_for_file(filename, url_map)
    solution_id = extract_solution_from_filename(filename)

    # Get all text from document
    full_text = get_paragraphs_text(doc)

    if not full_text or len(full_text) < 50:
        return []

    # Clean and truncate
    cleaned_text = clean_update_text(full_text)

    if not is_meaningful_update(cleaned_text):
        return []

    if len(cleaned_text) > MAX_UPDATE_LENGTH:
        cleaned_text = cleaned_text[:MAX_UPDATE_LENGTH] + '...[truncated]'

    return [{
        'update_id': generate_update_id(),
        'solution_id': normalize_solution_id(solution_id),
        'update_text': cleaned_text,
        'source_document': 'SEP Meeting',
        'source_category': 'SEP',
        'source_url': source_url,
        'source_tab': filename,
        'meeting_date': meeting_date or '',
        'created_at': datetime.now().isoformat(),
        'created_by': 'sep_import'
    }]


def parse_consolidated_document(doc_path, url_map):
    """Parse consolidated SEP document with multiple meeting dates"""
    try:
        doc = docx.Document(doc_path)
    except Exception as e:
        print(f"  Error opening {doc_path.name}: {e}")
        return []

    filename = doc_path.name
    source_url = find_url_for_file(filename, url_map)

    # Extract year from filename
    year_match = re.search(r'_C0_(\d{4})\.docx', filename)
    current_year = int(year_match.group(1)) if year_match else datetime.now().year

    updates = []
    current_meeting_date = None
    current_content = []
    date_pattern = re.compile(r'^(\d{2})_(\d{2})$')

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Check for date marker
        date_match = date_pattern.match(text)
        if date_match:
            # Save previous meeting's content
            if current_meeting_date and current_content:
                combined = '\n'.join(current_content)
                combined = clean_update_text(combined)

                if is_meaningful_update(combined):
                    if len(combined) > MAX_UPDATE_LENGTH:
                        combined = combined[:MAX_UPDATE_LENGTH] + '...[truncated]'

                    updates.append({
                        'update_id': generate_update_id(),
                        'solution_id': 'SEP',
                        'update_text': combined,
                        'source_document': 'SEP Weekly Meeting',
                        'source_category': 'SEP',
                        'source_url': source_url,
                        'source_tab': current_meeting_date,
                        'meeting_date': current_meeting_date,
                        'created_at': datetime.now().isoformat(),
                        'created_by': 'sep_consolidated_import'
                    })

            # Start new meeting
            month, day = date_match.groups()
            current_meeting_date = f"{current_year}-{month}-{day}"
            current_content = []
        else:
            current_content.append(text)

    # Don't forget last meeting
    if current_meeting_date and current_content:
        combined = '\n'.join(current_content)
        combined = clean_update_text(combined)

        if is_meaningful_update(combined):
            if len(combined) > MAX_UPDATE_LENGTH:
                combined = combined[:MAX_UPDATE_LENGTH] + '...[truncated]'

            updates.append({
                'update_id': generate_update_id(),
                'solution_id': 'SEP',
                'update_text': combined,
                'source_document': 'SEP Weekly Meeting',
                'source_category': 'SEP',
                'source_url': source_url,
                'source_tab': current_meeting_date,
                'meeting_date': current_meeting_date,
                'created_at': datetime.now().isoformat(),
                'created_by': 'sep_consolidated_import'
            })

    return updates


def main():
    all_updates = []
    files_processed = 0

    print("Building URL mapping from file log...")
    url_map = build_url_mapping()
    print(f"  Loaded {len(url_map)} URL mappings")
    print()

    print("Extracting SEP updates from Word documents...")
    print(f"Base path: {BASE_PATH}")
    print()

    # Process individual files in root SEP folder
    root_files = list(BASE_PATH.glob("*.docx"))
    if root_files:
        print(f"Processing root folder ({len(root_files)} files)...")
        for doc_file in sorted(root_files):
            updates = parse_document(doc_file, url_map)
            if updates:
                all_updates.extend(updates)
            files_processed += 1
        print(f"  Found {len(all_updates)} updates")

    # Process SEP Weekly Meeting Notes
    weekly_path = BASE_PATH / "SEP - SNWG Weekly Meeting Notes"
    if weekly_path.exists():
        # FY subfolders
        for fy_folder in sorted(weekly_path.glob("FY*")):
            if fy_folder.is_dir():
                docx_files = list(fy_folder.glob("*.docx"))
                print(f"Processing {fy_folder.name} ({len(docx_files)} files)...")
                fy_count = 0
                for doc_file in sorted(docx_files):
                    updates = parse_document(doc_file, url_map)
                    if updates:
                        fy_count += len(updates)
                        all_updates.extend(updates)
                    files_processed += 1
                print(f"  Found {fy_count} updates")

        # Consolidated files
        for doc_file in sorted(weekly_path.glob("*_C0_*.docx")):
            print(f"Processing consolidated: {doc_file.name}")
            updates = parse_consolidated_document(doc_file, url_map)
            if updates:
                print(f"  Found {len(updates)} updates")
                all_updates.extend(updates)
            files_processed += 1

    # Process SEP OPERA folder
    opera_path = BASE_PATH / "SEP OPERA"
    if opera_path.exists():
        docx_files = list(opera_path.glob("*.docx"))
        if docx_files:
            print(f"Processing SEP OPERA ({len(docx_files)} files)...")
            opera_count = 0
            for doc_file in sorted(docx_files):
                updates = parse_document(doc_file, url_map)
                if updates:
                    # Override solution_id for OPERA files
                    for u in updates:
                        u['solution_id'] = 'OPERA'
                    opera_count += len(updates)
                    all_updates.extend(updates)
                files_processed += 1
            print(f"  Found {opera_count} updates")

    # Process SEP Cycle 3 (SPoRT)
    sport_path = BASE_PATH / "SEP Cycle 3 (SPoRT)"
    if sport_path.exists():
        docx_files = list(sport_path.glob("*.docx"))
        if docx_files:
            print(f"Processing SEP SPoRT ({len(docx_files)} files)...")
            sport_count = 0
            for doc_file in sorted(docx_files):
                updates = parse_document(doc_file, url_map)
                if updates:
                    sport_count += len(updates)
                    all_updates.extend(updates)
                files_processed += 1
            print(f"  Found {sport_count} updates")

    print()
    print("=" * 60)
    print(f"Total files processed: {files_processed}")
    print(f"Total updates found: {len(all_updates)}")

    # Count by solution
    by_solution = {}
    for u in all_updates:
        sol = u['solution_id']
        by_solution[sol] = by_solution.get(sol, 0) + 1

    print(f"Unique solutions: {len(by_solution)}")
    print("\nUpdates by solution:")
    for sol, count in sorted(by_solution.items(), key=lambda x: -x[1]):
        print(f"  {sol}: {count}")

    # Count with URLs
    with_urls = sum(1 for u in all_updates if u.get('source_url'))
    print(f"\nUpdates with source URLs: {with_urls}/{len(all_updates)}")

    # Write Excel
    if all_updates:
        OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)

        df = pd.DataFrame(all_updates, columns=OUTPUT_COLUMNS)

        # Sort by date descending
        df['meeting_date'] = pd.to_datetime(df['meeting_date'], errors='coerce')
        df = df.sort_values(['meeting_date', 'solution_id'], ascending=[False, True])
        df['meeting_date'] = df['meeting_date'].dt.strftime('%Y-%m-%d')

        # Remove rows without dates
        df = df[df['meeting_date'].notna()]

        with pd.ExcelWriter(OUTPUT_PATH, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='SEP Updates', index=False)

            worksheet = writer.sheets['SEP Updates']
            for idx, col in enumerate(df.columns):
                max_length = max(df[col].astype(str).map(len).max(), len(col))
                adjusted_width = min(max_length + 2, 60)
                col_letter = chr(65 + idx) if idx < 26 else f"A{chr(65 + idx - 26)}"
                worksheet.column_dimensions[col_letter].width = adjusted_width

        print(f"\nExcel file written to: {OUTPUT_PATH}")
    else:
        print("\nNo updates found to export.")


if __name__ == '__main__':
    main()
